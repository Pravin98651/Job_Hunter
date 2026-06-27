import asyncio
import ipaddress
import logging
import os
import shutil
import socket
import tempfile
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from playwright.async_api import async_playwright, Page

logger = logging.getLogger(__name__)


def _validate_apply_url(url: str) -> None:
    """
    Validate that the apply URL is not pointing to a private/loopback/reserved
    address (SSRF prevention). Raises ValueError if invalid.
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"Invalid URL scheme: {parsed.scheme}")
    hostname = parsed.hostname
    if not hostname:
        raise ValueError("URL has no hostname")
    try:
        resolved_ips = socket.getaddrinfo(hostname, None)
    except socket.gaierror as e:
        raise ValueError(f"Cannot resolve hostname '{hostname}': {e}")
    for _, _, _, _, sockaddr in resolved_ips:
        ip_str = sockaddr[0]
        try:
            ip = ipaddress.ip_address(ip_str)
            if ip.is_private or ip.is_loopback or ip.is_reserved or ip.is_link_local:
                raise ValueError(f"URL resolves to restricted IP: {ip_str}")
        except ValueError as e:
            raise ValueError(f"Invalid resolved IP: {ip_str} — {e}")

async def fill_heuristic_fields(page: Page, profile: dict):
    """
    Attempts to map the user's profile data to standard application form fields.
    """
    # Simple dictionary of regex patterns to data values
    
    # Get basic info from profile
    first_name = ""
    last_name = ""
    name_parts = profile.get("name", "").split(" ")
    if len(name_parts) >= 2:
        first_name = name_parts[0]
        last_name = " ".join(name_parts[1:])
    else:
        first_name = profile.get("name", "")

    mappings = [
        ({"name": "first", "placeholder": "first"}, first_name),
        ({"name": "last", "placeholder": "last"}, last_name),
        ({"name": "email", "placeholder": "email", "type": "email"}, profile.get("email", "")),
        ({"name": "phone", "placeholder": "phone", "type": "tel"}, profile.get("phone", "")),
        ({"name": "linkedin", "placeholder": "linkedin"}, profile.get("linkedin_url", "")),
        ({"name": "github", "placeholder": "github"}, profile.get("github_url", "")),
        ({"name": "portfolio", "placeholder": "website"}, profile.get("portfolio_url", ""))
    ]

    for mapping, value in mappings:
        if not value:
            continue
        
        # Build a flexible selector string. Playwright allows matching multiple.
        # e.g. input[name*="first" i], input[placeholder*="first" i]
        selectors = []
        if "name" in mapping:
            selectors.append(f'input[name*="{mapping["name"]}" i]')
            selectors.append(f'input[id*="{mapping["name"]}" i]')
        if "placeholder" in mapping:
            selectors.append(f'input[placeholder*="{mapping["placeholder"]}" i]')
        if "type" in mapping:
            selectors.append(f'input[type="{mapping["type"]}"]')

        selector = ", ".join(selectors)
        try:
            # Look for the elements
            elements = await page.locator(selector).all()
            for el in elements:
                is_visible = await el.is_visible()
                is_enabled = await el.is_enabled()
                if is_visible and is_enabled:
                    current_val = await el.input_value()
                    if not current_val:
                        await el.fill(value)
        except Exception:
            pass


async def run_auto_fill(apply_url: str, resume_profile: dict, cover_letter_text: str | None = None, resume_bytes: bytes | None = None, headless: bool = False) -> bool:
    """
    Launches a chromium browser, navigates to apply_url, attempts heuristic 
    filling. If headless is True, it will attempt to autonomously click Submit/Apply and return True/False based on success. If False, it waits for the user to close the browser manually.
    """
    if not apply_url or not apply_url.startswith("http"):
        logger.warning(f"[auto_apply] Invalid apply_url provided: {apply_url!r}")
        return False

    # SSRF validation — prevent the browser from navigating to internal services
    try:
        _validate_apply_url(apply_url)
    except ValueError as e:
        logger.warning(f"[auto_apply] Blocked unsafe URL: {e}")
        return False

    # Create temporary cover letter file if text is provided
    temp_cover_letter_path = None
    if cover_letter_text:
        doc = Document()
        doc.add_heading(f"Cover Letter", 0)
        doc.add_paragraph(cover_letter_text)
        
        # Save to temp
        fd, temp_cover_letter_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(temp_cover_letter_path)

    # Create temporary resume file if bytes are provided
    temp_resume_path = None
    if resume_bytes:
        fd, temp_resume_path = tempfile.mkstemp(suffix=".pdf")
        with os.fdopen(fd, 'wb') as f:
            f.write(resume_bytes)

    temp_user_data_dir = tempfile.mkdtemp(prefix="playwright_profile_")
    
    async with async_playwright() as p:
        args = [
            "--window-size=1280,800",
            "--disable-blink-features=AutomationControlled",
            "--disable-extensions",
            "--disable-default-apps",
            "--disable-sync",
            "--no-first-run",
            "--incognito",
            "--disable-popup-blocking=false", # Force popup blocking
        ]
        
        context = await p.chromium.launch_persistent_context(
            user_data_dir=temp_user_data_dir,
            headless=headless,
            args=args,
            viewport={"width": 1280, "height": 800},
            accept_downloads=False
        )
        page = context.pages[0] if context.pages else await context.new_page()

        try:
            await page.goto(apply_url, wait_until="domcontentloaded", timeout=60000)
            
            # Wait a few seconds for single-page apps / dynamic forms to load
            await page.wait_for_timeout(3000)

            # Heuristic text field filling
            if resume_profile:
                await fill_heuristic_fields(page, resume_profile)

            # Heuristic file uploads
            # Find inputs of type "file"
            file_inputs = await page.locator('input[type="file"]').all()
            for f_input in file_inputs:
                try:
                    # check label or name context
                    id_val = await f_input.get_attribute("id") or ""
                    name_val = await f_input.get_attribute("name") or ""
                    
                    # Heuristically check if it's looking for a resume or cover letter
                    is_resume = "resume" in id_val.lower() or "resume" in name_val.lower() or "cv" in id_val.lower() or "cv" in name_val.lower()
                    is_cover = "cover" in id_val.lower() or "cover" in name_val.lower() or "letter" in id_val.lower()

                    if is_resume and temp_resume_path:
                        await f_input.set_input_files(temp_resume_path)
                    elif is_cover and temp_cover_letter_path:
                        await f_input.set_input_files(temp_cover_letter_path)
                    elif not is_resume and not is_cover:
                        # Fallback: if there's only one file input, usually it's the resume
                        if len(file_inputs) == 1 and temp_resume_path:
                            await f_input.set_input_files(temp_resume_path)
                except Exception:
                    pass

            if headless:
                # Autonomous mode: attempt to submit the application
                try:
                    submit_selectors = [
                        "button[type='submit']",
                        "input[type='submit']",
                        "button:has-text('Submit')",
                        "button:has-text('Apply')",
                        "button:has-text('Next')"
                    ]
                    clicked = False
                    for sel in submit_selectors:
                        if clicked: break
                        buttons = await page.locator(sel).all()
                        for btn in buttons:
                            if await btn.is_visible() and await btn.is_enabled():
                                await btn.click()
                                await page.wait_for_timeout(3000)
                                clicked = True
                                break
                    return True
                except Exception as e:
                    logger.error(f"[auto_apply] Autonomous submission error: {e}")
                    return False
            else:
                # Interactive mode: Wait for user to finish and close
                try:
                    await page.wait_for_event("close", timeout=300000)
                except Exception:
                    logger.info("[auto_apply] Playwright session timed out or user closed the browser")
                return True
            
        except Exception as e:
            logger.error(f"[auto_apply] Playwright automation error: {e}", exc_info=True)
            return False
        finally:
            await context.close()

    # Cleanup temp files outside the playwright context to ensure they're always cleaned
    try:
        shutil.rmtree(temp_user_data_dir, ignore_errors=True)
    except Exception:
        pass
    if temp_cover_letter_path and os.path.exists(temp_cover_letter_path):
        os.remove(temp_cover_letter_path)
    if temp_resume_path and os.path.exists(temp_resume_path):
        os.remove(temp_resume_path)
